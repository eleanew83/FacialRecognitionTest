#!/usr/bin/env python
"""Generate an academic-style Word report summarising the SAM 3 experiments."""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUT_PATH = Path(
    "/rds/user/ylj20/hpc-work/FacialRecognitionTest/sam3_experiments/"
    "SAM3_macaque_detection_report.docx"
)

# ── Document setup ────────────────────────────────────────────────────────────

doc = Document()

# Set default font to Times New Roman 11pt (closer to journal style)
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)

# Set page margins
for section in doc.sections:
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, italic=False, bold=False, align=None, size=11):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.italic = italic
    run.bold   = bold
    return p


def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.italic = True


def add_table(headers, rows, highlight_row=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"

    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
            if highlight_row is not None and r_idx == highlight_row:
                run.bold = True
    return table


# ── Title ─────────────────────────────────────────────────────────────────────

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
t_run = title.add_run(
    "Zero-Shot Macaque Face Detection with SAM 3: "
    "A Prompt-Benchmark and Comparison Against a Trained YOLO Detector"
)
t_run.font.name = "Times New Roman"
t_run.font.size = Pt(14)
t_run.bold = True

add_para(
    f"FacialRecognitionTest project — sam3_experiments/  "
    f"({datetime.now().strftime('%Y-%m-%d')})",
    italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10,
)

# ── Abstract ──────────────────────────────────────────────────────────────────

add_heading("Abstract", level=1)
add_para(
    "We evaluate the zero-shot, text-promptable Segment Anything Model 3 "
    "(SAM 3) on the task of macaque face detection in field photographs of "
    "the Gibraltar Barbary macaque (Macaca sylvanus) population, and compare "
    "it against an in-house YOLOv8 face detector trained on the same dataset. "
    "Two complementary experiments were run on an NVIDIA A100 80 GB GPU: "
    "(i) a quantitative prompt benchmark against ground-truth bounding boxes "
    "on a 50-image stratified sample of the YOLO validation split, sweeping "
    "13 candidate text prompts under a common confidence threshold "
    "(score > 0.05) and non-maximum suppression pipeline (IoU > 0.5); and "
    "(ii) a qualitative three-scenario test (single face, multi-face, "
    "macaque-with-human) on real Gibraltar and Instagram tourist photographs. "
    "The trained YOLO detector achieved AP50 = 0.981 on the full validation "
    "set (n = 1,333; 0.979 on the 50-image sample), with precision = recall "
    "= F1 = 0.98. The best SAM 3 prompt, “monkey face”, reached AP50 = 0.621 "
    "(precision 0.57, recall 0.78, F1 = 0.66), 0.358 AP50 below the trained "
    "detector. Counter-intuitively, species-specific prompts (“macaque face”, "
    "“Barbary macaque face”) underperformed the generic taxon-level prompt, "
    "consistent with the relative frequencies of these terms in CLIP-style "
    "pre-training corpora. In Scenario 3 (macaque + human), SAM 3 recovered "
    "human faces that the macaque-specific YOLO model is incapable of "
    "detecting in 4 / 10 images, demonstrating a complementary role for SAM 3 "
    "as an open-vocabulary detector. We conclude that a trained YOLO model "
    "remains the operational detector of choice for the existing macaque "
    "re-identification pipeline, while SAM 3 is best deployed for "
    "open-vocabulary tasks (human-presence filtering, novel-prompt "
    "exploration) where labelled data are unavailable.",
)

# ── 1. Introduction ───────────────────────────────────────────────────────────

add_heading("1. Introduction", level=1)
add_para(
    "Reliable face localisation is the first stage of any individual "
    "re-identification pipeline. In the FacialRecognitionTest project we "
    "have built and trained a YOLOv8 macaque-face detector that powers the "
    "subsequent embedding and identification stages. Meta’s release of "
    "Segment Anything Model 3 (SAM 3), which accepts a free-form text prompt "
    "and returns segmentation masks plus bounding boxes, raises the natural "
    "question of whether a general-purpose, zero-shot model can replace or "
    "supplement the trained detector. SAM 3 also opens the door to "
    "operations that are out of scope for a species-specific YOLO model, "
    "most notably the detection of human faces in the same frame as "
    "macaques — an important downstream consideration for tourist-photo "
    "data sources (Instagram, Cable Car location)."
)
add_para(
    "This report summarises the experiments conducted in "
    "sam3_experiments/ during February 2026. We address three questions:"
)
add_para(
    "  (Q1) Which textual prompt elicits the strongest macaque-face detector "
    "from SAM 3?\n"
    "  (Q2) How does the best SAM 3 prompt compare quantitatively to the "
    "trained YOLO baseline?\n"
    "  (Q3) Does SAM 3 expand the operational envelope of the existing "
    "pipeline (multi-face frames, mixed macaque/human scenes)?"
)

# ── 2. Materials and Methods ──────────────────────────────────────────────────

add_heading("2. Materials and Methods", level=1)

add_heading("2.1 Hardware and software", level=2)
add_para(
    "All experiments were run on a single NVIDIA A100-SXM4-80 GB GPU "
    "(node gpu-q-2 / gpu-q-42), with PyTorch 2.7.0 + CUDA 12.6, Python "
    "3.12 inside a dedicated conda environment (sam3). The SAM 3 model "
    "weights were obtained from the facebook/sam3 HuggingFace repository "
    "(gated access). SAM 3 does not currently support CPU inference; all "
    "code paths assume CUDA. The YOLO baseline used the in-house checkpoint "
    "macaque_face_detector_20260120_v1/best.pt loaded via ultralytics 8.4.14."
)

add_heading("2.2 Datasets", level=2)
add_para(
    "Two image sources were used. (i) The Gibraltar Macaques dataset "
    "comprises approximately 9,900 field photographs organised by location "
    "and named individual (folders females/ and males/). For the "
    "ground-truth benchmark, the existing YOLO validation split was used "
    "(1,333 images, one face per image as the dominant case; mean face count "
    "1.0). A seeded sample of 50 images was drawn from this split for the "
    "prompt sweep. (ii) The IG_Photos dataset (n = 10) contains real "
    "Instagram-scraped tourist photographs that include both human and "
    "macaque faces in the same frame, and is used exclusively for the "
    "Scenario 3 qualitative test."
)

add_heading("2.3 Model invocation and post-processing", level=2)
add_para(
    "Each input image was passed to SAM 3 in image-mode (Sam3Processor) with "
    "a single text prompt per inference call. The raw output (boxes, scores, "
    "masks) was passed through a two-step filter:"
)
add_para(
    "  (a) confidence threshold: boxes with score ≤ 0.05 were discarded;\n"
    "  (b) non-maximum suppression at IoU > 0.5 (with an additional "
    "intersection-over-min-area rule of 0.5 to catch the case where a tight "
    "box is nested inside a loose box around the same face)."
)
add_para(
    "These same thresholds were applied identically to every prompt so the "
    "comparison is fair. For YOLO, the standard confidence threshold of 0.25 "
    "was used (Ultralytics default)."
)

add_heading("2.4 Prompts evaluated", level=2)
add_para(
    "Thirteen prompts were swept, spanning four families: species-specific "
    "(“macaque face”, “Barbary macaque face”, “Barbary macaque”), taxon-level "
    "(“monkey face”, “monkey head”, “macaque head”, “primate face”), "
    "visual/descriptive (“furry animal face”, “close-up of a monkey face”, "
    "“animal face”), body-level (“macaque”, “monkey”), and a zero-shot "
    "ceiling (“face”)."
)

add_heading("2.5 Evaluation metrics", level=2)
add_para(
    "Predicted boxes were matched to ground-truth boxes by greedy "
    "highest-IoU assignment; a prediction was a true positive if its best "
    "IoU with an unmatched GT box was ≥ 0.5. We report precision = "
    "TP / (TP + FP), recall = TP / (TP + FN), F1, and AP50 (area under the "
    "precision–recall curve at IoU 0.5). Per-image AP50 was also recorded; "
    "the dataset-level AP50 reported here is the mean of per-image values "
    "as written by benchmark_prompts.py."
)

# ── 3. Experiments ────────────────────────────────────────────────────────────

add_heading("3. Experiments", level=1)

add_heading("3.1 Quantitative prompt benchmark (Experiment A)", level=2)
add_para(
    "Each of the 13 prompts was applied to every image in the 50-image "
    "validation sample; predicted boxes were filtered and NMS-deduplicated, "
    "and matched against the YOLO-format ground truth labels at IoU ≥ 0.5. "
    "Per-image side-by-side visualisations (GT vs. SAM 3 vs. YOLO predictions) "
    "are saved under results/eval_vs_yolo/ together with the summary CSV. "
    "For reference, the trained YOLO model was first evaluated on the full "
    "validation set (n = 1,333) using its own val() routine to anchor the "
    "numbers reported on the 50-image sample."
)

add_heading("3.2 Qualitative scenario test (Experiment B)", level=2)
add_para(
    "A complementary three-scenario test was run on raw photographs to "
    "characterise SAM 3’s behaviour outside the labelled benchmark."
)
add_para(
    "  S1 — Single macaque face: three images sampled (seed = 42) from "
    "individual folders, with no “+” character in the filename (single "
    "named animal).\n"
    "  S2 — Multiple macaque faces: three images whose filenames contain "
    "“+”, indicating two or more named animals (23 such images exist in "
    "the corpus).\n"
    "  S3 — Macaque + human face: all 10 IG_Photos images; for each frame, "
    "SAM 3 was run twice (prompts “monkey face” and “human face”) and the "
    "union plotted in two colours."
)
add_para(
    "Throughout Scenario 3 the SAM 3 prompt “monkey face” was used as it was "
    "the top-ranked prompt from Experiment A."
)

# ── 4. Results ────────────────────────────────────────────────────────────────

add_heading("4. Results", level=1)

add_heading("4.1 YOLO baseline", level=2)
add_para(
    "On the full validation set (1,333 images, 1,333 instances), the trained "
    "YOLO detector achieved AP50 = 0.981 and AP50-95 = 0.600, with "
    "P = 0.976, R = 0.964 (Ultralytics val()). On the 50-image sample, the "
    "matched metrics were AP50 = 0.979, P = 0.98, R = 0.98, F1 = 0.98, "
    "TP = 49, FP = 1, FN = 1, confirming that the sample is representative "
    "of the full split (sample AP50 within 0.002 of full)."
)

add_heading("4.2 Prompt benchmark", level=2)
add_para(
    "Table 1 reports the 13 prompts ranked by AP50, with the YOLO reference "
    "row at the top. Counts are TP/FP/FN against 50 ground-truth boxes "
    "(one per image)."
)
add_caption(
    "Table 1. SAM 3 prompt benchmark on the 50-image validation sample "
    "(IoU ≥ 0.5). YOLO (trained) is the in-house macaque face detector "
    "baseline. Bold = best SAM 3 prompt."
)

headers = ["Rank", "Detector / Prompt", "AP50", "Precision", "Recall",
           "F1", "TP", "FP", "FN"]
rows = [
    ["REF", "YOLO (trained)",              "0.979", "0.980", "0.980", "0.980", 49,  1,  1],
    ["1",   "SAM3: monkey face",           "0.621", "0.565", "0.780", "0.655", 39, 30, 11],
    ["2",   "SAM3: monkey head",           "0.585", "0.708", "0.680", "0.694", 34, 14, 16],
    ["3",   "SAM3: macaque head",          "0.477", "0.690", "0.580", "0.630", 29, 13, 21],
    ["4",   "SAM3: face",                  "0.452", "0.589", "0.660", "0.623", 33, 23, 17],
    ["5",   "SAM3: close-up of a monkey face","0.422","0.397","0.620","0.484", 31, 47, 19],
    ["6",   "SAM3: macaque face",          "0.346", "0.587", "0.540", "0.563", 27, 19, 23],
    ["7",   "SAM3: primate face",          "0.269", "0.682", "0.300", "0.417", 15,  7, 35],
    ["8",   "SAM3: Barbary macaque face",  "0.238", "0.700", "0.280", "0.400", 14,  6, 36],
    ["9",   "SAM3: furry animal face",     "0.063", "0.571", "0.080", "0.140",  4,  3, 46],
    ["10",  "SAM3: animal face",           "0.040", "1.000", "0.040", "0.077",  2,  0, 48],
    ["11",  "SAM3: monkey",                "0.004", "0.051", "0.060", "0.055",  3, 56, 47],
    ["12",  "SAM3: Barbary macaque",       "0.003", "0.038", "0.040", "0.039",  2, 50, 48],
    ["13",  "SAM3: macaque",               "0.001", "0.026", "0.020", "0.023",  1, 37, 49],
]
add_table(headers, rows, highlight_row=2)  # highlight "monkey face" row

add_para("")
add_para(
    "Three findings stand out. First, the best SAM 3 prompt (“monkey face”) "
    "trails the trained YOLO detector by 0.358 AP50 (0.621 vs. 0.979) and "
    "by 0.32 F1 (0.66 vs. 0.98) on the same images. Second, taxon-level "
    "prompts (“monkey *”, AP50 0.585–0.621) consistently outperform "
    "species-specific prompts (“macaque face” 0.346; “Barbary macaque face” "
    "0.238). Third, dropping the “face” / “head” qualifier collapses "
    "performance: “monkey” alone scores AP50 = 0.004 (FP = 56) because SAM "
    "3 then returns whole-body boxes that fail the IoU = 0.5 test against "
    "face-only GT. The bias is therefore concentrated at two axes: lexical "
    "frequency in pre-training data (monkey ≫ Barbary macaque) and the "
    "presence of an explicit anatomical anchor (“face”/“head”)."
)

add_heading("4.3 Scenario test", level=2)
add_para(
    "Table 2 reports the per-image detection counts in the three qualitative "
    "scenarios, using the best prompt “monkey face” (and additionally "
    "“human face” in S3)."
)
add_caption(
    "Table 2. Per-image detection counts for the qualitative scenarios. "
    "S3 reports SAM 3 results decomposed by prompt; the “SAM 3 total” "
    "column is the union over both prompts after independent NMS per prompt."
)

headers2 = ["Scenario", "Image", "YOLO", "SAM 3 (monkey face)",
            "SAM 3 (human face)", "SAM 3 total"]
rows2 = [
    ["S1", "Abby (Apes Den)",                              2, 0, "—", 0],
    ["S1", "Jordi (Prince Philip Arch)",                   1, 1, "—", 1],
    ["S1", "Jamie temporary (O’Hara)",                     1, 1, "—", 1],
    ["S2", "Tripatte + AF2 (Rock Gun)",                    1, 0, "—", 0],
    ["S2", "Flick + follow (O’Hara)",                      1, 1, "—", 1],
    ["S2", "Alexa + SAM (O’Hara)",                         2, 2, "—", 2],
    ["S3", "image_10",                                      1, 2,   1, 3],
    ["S3", "image_4",                                       2, 2,   0, 2],
    ["S3", "image_8",                                       1, 2,   1, 3],
    ["S3", "image_44",                                      1, 1,   1, 2],
    ["S3", "image_33",                                      2, 0,   1, 1],
    ["S3", "image_5",                                       1, 1,   0, 1],
    ["S3", "image_28",                                      1, 1,   1, 2],
    ["S3", "image_2",                                       2, 2,   0, 2],
    ["S3", "image_18",                                      2, 2,   0, 2],
    ["S3", "image_3",                                       1, 1,   0, 1],
]
add_table(headers2, rows2)

add_para("")
add_para(
    "Aggregated over the 10 IG_Photos frames in S3, YOLO returned 14 "
    "macaque detections, SAM 3 returned 14 macaque detections plus 5 "
    "human-face detections that YOLO cannot, by construction, produce. "
    "SAM 3 exceeded YOLO’s box count in 4 of 10 S3 frames (image_10, "
    "image_8, image_44, image_28) and was at parity on 5; on image_33 SAM "
    "3 missed both macaques while still recovering a human face. In S1/S2, "
    "SAM 3 matched YOLO on 4 of 6 frames and missed faces present in YOLO "
    "output on 2 (Abby, Tripatte) — both are field photographs in difficult "
    "lighting and partial occlusion conditions."
)

# ── 5. Discussion ─────────────────────────────────────────────────────────────

add_heading("5. Discussion", level=1)

add_heading("5.1 Why generic prompts beat species-specific ones", level=2)
add_para(
    "The 0.275 AP50 gap between “monkey face” and “macaque face”, and the "
    "0.383 gap between “monkey face” and “Barbary macaque face”, are too "
    "large to attribute to noise on a 50-image sample. The most plausible "
    "explanation is the underlying distribution of these terms in the "
    "image–text corpora used to pre-train SAM 3’s text encoder: web "
    "captions contain orders of magnitude more occurrences of “monkey” "
    "than of “macaque”, let alone “Barbary macaque”. The encoder therefore "
    "produces a stronger and better-calibrated visual concept for the "
    "common term. This is consistent with prior reports on CLIP-style "
    "encoders, where rare-class taxonomic names lag behind common-language "
    "labels."
)

add_heading("5.2 Anatomical anchors matter", level=2)
add_para(
    "Removing the word “face” or “head” from a prompt drops AP50 by two "
    "orders of magnitude (0.621 → 0.004 for “monkey face” → “monkey”). "
    "SAM 3 then returns whole-animal boxes that do not satisfy the "
    "IoU ≥ 0.5 condition against face-only ground truth. For pipelines "
    "downstream of detection (face embedding, individual re-identification) "
    "this matters: a body-box is not interchangeable with a face-box. Any "
    "SAM 3 deployment for face detection should include an explicit "
    "anatomical anchor in the prompt."
)

add_heading("5.3 Where SAM 3 adds value over the trained detector", level=2)
add_para(
    "The Scenario 3 results expose a capability the trained YOLO model "
    "does not have. The macaque face detector was trained on macaque faces "
    "only and cannot, by design, flag human faces in the same frame. In "
    "the IG_Photos dataset, where tourist selfies routinely include both "
    "humans and macaques, SAM 3’s zero-shot “human face” prompt recovered "
    "five human-face detections across the ten frames at the same "
    "confidence and NMS thresholds. This is directly useful: it offers a "
    "drop-in pre-filter to either exclude human-containing frames from "
    "the macaque re-identification pipeline or to flag them for manual "
    "review, without retraining a multi-class detector."
)

add_heading("5.4 Limitations", level=2)
add_para(
    "Three limitations should be flagged. (i) Sample size: the 50-image "
    "benchmark is a representative sample (YOLO AP50 within 0.002 of the "
    "full split) but a 200–500 image evaluation would tighten the prompt "
    "ranking, especially in the middle of the table where AP50 differences "
    "are within 0.05–0.10. (ii) Single-prompt inference: each SAM 3 call "
    "uses one prompt; an ensemble of the top 2–3 prompts followed by "
    "cross-prompt NMS was not tested and is likely to improve recall. "
    "(iii) Confidence threshold: a fixed score > 0.05 was used to keep "
    "the prompt comparison fair, but for downstream use the threshold "
    "should be re-tuned per prompt against an operating-point preference "
    "(precision vs. recall)."
)

# ── 6. Conclusion ─────────────────────────────────────────────────────────────

add_heading("6. Conclusion", level=1)
add_para(
    "On macaque face detection in the FacialRecognitionTest corpus, a "
    "trained YOLO detector remains substantially ahead of zero-shot SAM 3 "
    "(AP50 0.979 vs. 0.621; F1 0.98 vs. 0.66). However, SAM 3 is not "
    "redundant: it provides a low-effort, open-vocabulary tool for tasks "
    "the macaque-specific YOLO cannot address — most importantly the "
    "detection of human faces in mixed scenes. We recommend keeping the "
    "trained YOLO detector as the primary face localiser in the "
    "re-identification pipeline, and using SAM 3 with the prompt “monkey "
    "face” (or an ensemble with “monkey head”) as an auxiliary detector for "
    "(a) frames containing human bystanders, (b) prompts beyond the YOLO "
    "training classes, and (c) data-curation passes where labelled data are "
    "not yet available."
)

# ── Appendix ──────────────────────────────────────────────────────────────────

add_heading("Appendix A. Files and reproducibility", level=1)
add_para(
    "All scripts and outputs live under sam3_experiments/."
)
add_para(
    "  • benchmark_prompts.py — prompt sweep and ground-truth evaluation "
    "(--eval mode)\n"
    "  • test_sam3_macaque.py — three-scenario qualitative test\n"
    "  • run_sam3_gpu.sh — SLURM submission wrapper (A100, conda env sam3)\n"
    "  • results/eval_vs_yolo/eval_vs_yolo.csv — prompt benchmark results\n"
    "  • results/eval_vs_yolo/eval_*.jpg — 50 side-by-side comparisons "
    "(GT, YOLO, SAM 3 best)\n"
    "  • results/scenario_20260225_044910/ — final scenario run, 16 panels "
    "(S1×3, S2×3, S3×10)\n"
    "  • logs/23352722.out — full eval log\n"
    "  • logs/23434961.out — final scenario log"
)
add_para(
    "Random seed = 42 throughout. SAM 3 inference is deterministic at the "
    "image level given a fixed prompt and post-processing pipeline."
)

# ── Save ──────────────────────────────────────────────────────────────────────

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
